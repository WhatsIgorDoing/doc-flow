"""
Script para criar fixtures de teste ausentes para o SAD_APP v2.0
"""

import openpyxl
from pathlib import Path
from docx import Document


def create_manifesto_exemplo():
    """Cria o arquivo manifesto_exemplo.xlsx"""
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Manifesto"
    
    # Cabeçalho
    ws['A1'] = 'document_code'
    ws['B1'] = 'revision'
    ws['C1'] = 'title'
    ws['D1'] = 'DISCIPLINA'
    
    # Dados
    ws['A2'] = 'documento_pid'
    ws['B2'] = 'A'
    ws['C2'] = 'Documento PID de Teste'
    ws['D2'] = 'PROCESSO'
    
    ws['A3'] = 'documento_rir'
    ws['B3'] = 'B'
    ws['C3'] = 'Documento RIR de Teste'
    ws['D3'] = 'ELÉTRICA'
    
    # Salva
    fixture_path = Path('tests/fixtures/manifesto_exemplo.xlsx')
    fixture_path.parent.mkdir(parents=True, exist_ok=True)
    wb.save(fixture_path)
    print(f"✅ Criado: {fixture_path}")


def create_documento_rir_pdf():
    """Cria o arquivo documento_rir.pdf com texto contendo código"""
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import letter
    
    fixture_path = Path('tests/fixtures/documento_rir.pdf')
    fixture_path.parent.mkdir(parents=True, exist_ok=True)
    
    # Cria PDF com reportlab
    c = canvas.Canvas(str(fixture_path), pagesize=letter)
    c.setFont("Helvetica", 12)
    
    # Adiciona texto com código do documento esperado pelo teste
    c.drawString(100, 750, "REGISTRO DE INSPEÇÃO E RECEBIMENTO")
    c.drawString(100, 730, "")
    c.drawString(100, 710, "Código do Documento: CZ6_RNEST_U22_3.1.1.1_CVL_RIR_B-22026A")
    c.drawString(100, 690, "Revisão: A")
    c.drawString(100, 670, "")
    c.drawString(100, 650, "Este é um documento de teste para extração de código.")
    c.drawString(100, 630, "O código CZ6_RNEST_U22_3.1.1.1_CVL_RIR_B-22026A deve ser extraído.")
    
    c.save()
    print(f"✅ Criado: {fixture_path}")


def create_documento_rir_docx():
    """Cria o arquivo documento_rir.docx com texto contendo código"""
    fixture_path = Path('tests/fixtures/documento_rir.docx')
    fixture_path.parent.mkdir(parents=True, exist_ok=True)
    
    doc = Document()
    doc.add_heading('REGISTRO DE INSPEÇÃO E RECEBIMENTO', 0)
    doc.add_paragraph('')
    doc.add_paragraph('Código do Documento: CZ6_RNEST_U22_3.1.1.1_CVL_RIR_B-22026A')
    doc.add_paragraph('Revisão: A')
    doc.add_paragraph('')
    doc.add_paragraph('Este é um documento de teste para extração de código.')
    doc.add_paragraph('O código CZ6_RNEST_U22_3.1.1.1_CVL_RIR_B-22026A deve ser extraído.')
    
    doc.save(fixture_path)
    print(f"✅ Criado: {fixture_path}")


def create_template_exemplo():
    """Cria o arquivo template_exemplo.xlsx"""
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Template"
    
    # Cabeçalho do template conforme esperado pelo teste
    ws['A1'] = 'DOCUMENTO'
    ws['B1'] = 'REVISÃO'
    ws['C1'] = 'TÍTULO'
    ws['D1'] = 'ARQUIVO'
    ws['E1'] = 'STATUS'
    ws['F1'] = 'DISCIPLINA'
    ws['G1'] = 'TIPO'
    ws['H1'] = 'PROPÓSITO'
    
    # Template vazio (será preenchido pelo sistema)
    
    fixture_path = Path('tests/fixtures/template_exemplo.xlsx')
    fixture_path.parent.mkdir(parents=True, exist_ok=True)
    wb.save(fixture_path)
    print(f"✅ Criado: {fixture_path}")


def main():
    """Cria todos os fixtures necessários"""
    print("🔧 Criando fixtures de teste...\n")
    
    try:
        create_manifesto_exemplo()
        create_documento_rir_pdf()
        create_documento_rir_docx()
        create_template_exemplo()
        
        print("\n✅ Todos os fixtures foram criados com sucesso!")
        print("\n📋 Próximos passos:")
        print("1. Execute: python -m pytest tests/ -v")
        print("2. Verifique se os testes de integração passam")
        
    except Exception as e:
        print(f"\n❌ Erro ao criar fixtures: {e}")
        import traceback
        traceback.print_exc()


if __name__ == "__main__":
    main()
